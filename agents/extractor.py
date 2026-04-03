"""
agents/extractor.py
DITA Converter Tool — Extractor Agent

Parses PDF and DOCX files into a normalised Content Tree (list of block dicts).
Each block is produced by make_block() and carries: type, text, metadata.

Font-size thresholds calibrated against Gilbarco Passport manuals:
  H1 = 18pt  |  H2 = 14pt  |  H3 = 12pt
  Note header = 15pt  |  Steps/Figures = 10pt
  Body = 11pt  |  Headers/Footers = 9pt or less

Session: S-02 | Reviewer-signed-off
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# Block factory
# ---------------------------------------------------------------------------

VALID_TYPES = {
    "heading", "paragraph", "list_item", "table",
    "figure", "note_header", "note_inline", "code_block", "dropped",
}


def make_block(
    block_type: str,
    text: str,
    level: int = 0,
    is_header: bool = False,
    rows: list | None = None,
    metadata: dict | None = None,
) -> dict[str, Any]:
    if block_type not in VALID_TYPES:
        raise ValueError(f"Unknown block type: {block_type!r}")
    block: dict[str, Any] = {
        "type": block_type,
        "text": text.strip() if text else "",
        "level": level,
        "is_header": is_header,
        "rows": rows or [],
        "metadata": metadata or {},
        "dita_element": None,
    }
    return block


# ---------------------------------------------------------------------------
# Custom error
# ---------------------------------------------------------------------------

class ExtractorError(RuntimeError):
    pass


# ---------------------------------------------------------------------------
# Drop-pattern helpers
# ---------------------------------------------------------------------------

_DROP_PATTERNS = [
    re.compile(r"^Page \d+"),
    re.compile(r"MDE-\w+.+\d{4}$"),
    re.compile(r"^©\s*\d{4}"),
    # Note: "Table of Contents" and "Related Documents" are NOT dropped here —
    # they appear as section headings and also as table contexts on page 1.
    # The ROW_SHOW detector handles their table rows correctly.
]

# ---------------------------------------------------------------------------
# ROW_SHOW table detector
# ---------------------------------------------------------------------------
# FrameMaker ROW_SHOW tables have no vertical column lines.
# Headers are bounded by thick rules (~2pt, appears as rect height ≥ 1.5pt).
# Data rows are separated by thin rules (~0.5pt).
# Columns are inferred from word X-position clusters.
# ---------------------------------------------------------------------------

_ROW_SHOW_THICK = 1.5   # minimum rect height (pts) to be a header boundary rule
_ROW_SHOW_COL_GAP = 40  # fallback minimum — kept for straddle heuristics
_ROW_SHOW_MIN_GAP_PT = 7   # whitespace projection: min gap width (pts) to split columns
                           # With span-based coverage, intra-column inter-word gaps are
                           # typically 2-3pt, so 7pt safely separates real column gaps
                           # (which are >=8pt in observed MDE tables) from word spacing.


def _col_breaks_from_projection(words: list, x0: float, x1: float) -> list[float]:
    """
    Infer column boundaries from word spans using a filled-coverage projection.

    Builds a 1-pt boolean coverage array over the table width and marks every
    integer position covered by at least one word (x0..x1).  Contiguous zero
    zones (no word covers that x range) >= _ROW_SHOW_MIN_GAP_PT wide are true
    inter-column whitespace; a boundary is placed at each gap midpoint.

    Using spans (not just word starts) prevents false breaks caused by the
    intra-column inter-word gaps that occur when a multi-word column header
    (e.g. "UL Report Number") has words spread 8-30 pt apart.
    """
    width = int(x1 - x0) + 2
    density = [0] * width
    for w in words:
        xi0 = max(0, int(round(w["x0"] - x0)))
        xi1 = min(width - 1, int(round(w["x1"] - x0)))
        for xi in range(xi0, xi1 + 1):
            density[xi] = 1

    breaks: list[float] = [x0]
    in_gap = False
    gap_start = 0

    for i, occ in enumerate(density):
        if occ == 0 and not in_gap:
            in_gap, gap_start = True, i
        elif occ > 0 and in_gap:
            gap_w = i - gap_start
            if gap_w >= _ROW_SHOW_MIN_GAP_PT:
                mid = x0 + gap_start + gap_w / 2
                # Ignore tiny fragment breaks close to the previous boundary
                if mid - breaks[-1] > 15:
                    breaks.append(mid)
            in_gap = False

    breaks.append(x1 + 10)
    return breaks


def _extract_rowshow_tables(page) -> list[tuple[list[list[str]], float, float]]:
    """
    Detect and extract ROW_SHOW borderless tables from a pdfplumber page.

    Returns list of (rows, y_top) tuples.
    Each row is a list of cell strings. Multiple header rows supported.
    Straddled (spanning) cells are marked: cell[0]=text, cell[1]="__STRADDLE__{n_cols}"
    """
    from collections import defaultdict

    # Collect horizontal rule objects from both rects and lines.
    # Some PDFs (e.g. FrameMaker-generated MDE files) store table rules as
    # PDF line operators (page.lines) rather than filled rectangles.
    # Convert lines to synthetic rect-like dicts so the rest of the detector
    # can treat them uniformly.  Use linewidth to carry thickness information:
    # a line with linewidth >= _ROW_SHOW_THICK is considered a "thick" header
    # rule; thinner lines are row separators.
    rule_objects: list[dict] = list(page.rects)
    for ln in page.lines:
        if ln["x1"] - ln["x0"] > 50:  # horizontal span only
            lw = ln.get("linewidth", 0.5)
            rule_objects.append({
                **ln,
                "top":    ln["top"],
                "bottom": ln["top"] + lw,   # synthetic height = linewidth
            })
    rule_objects.sort(key=lambda r: r["top"])

    words = page.extract_words(
        x_tolerance=3, y_tolerance=5, extra_attrs=["fontname", "size"]
    )

    span_groups: dict = defaultdict(list)
    for r in rule_objects:
        if r["x1"] - r["x0"] > 50:
            key = (round(r["x0"], 0), round(r["x1"], 0))
            span_groups[key].append(r)

    results: list[tuple[list[list[str]], float, float]] = []

    for (x0, x1), group in span_groups.items():
        if len(group) < 3:
            continue

        group = sorted(group, key=lambda r: r["top"])

        # Split the group at large gaps: if two consecutive rules are >50pt
        # apart, they belong to separate tables that happen to share the same
        # x-span.  Treat each contiguous run as an independent group.
        sub_groups: list[list] = []
        cur: list = [group[0]]
        for r in group[1:]:
            if r["top"] - cur[-1]["bottom"] > 50:
                sub_groups.append(cur)
                cur = [r]
            else:
                cur.append(r)
        sub_groups.append(cur)

        for group in sub_groups:
            if len(group) < 3:
                continue

            thick = [r for r in group if (r["bottom"] - r["top"]) >= _ROW_SHOW_THICK]
            thin  = [r for r in group if (r["bottom"] - r["top"]) <  _ROW_SHOW_THICK]

            if len(thick) < 2:
                continue

            # Cluster consecutive thick rules into the "header cluster": a run of
            # adjacent thick rules where each is within 50pt of the previous one.
            # Thick rules further into the table are mid-table repeat-header rules;
            # treating all thick rules as header bands caused header_bottom to
            # encompass the entire table body, skipping all its data rows.
            header_cluster = [thick[0]]
            for r in thick[1:]:
                if r["top"] - header_cluster[-1]["bottom"] <= 50:
                    header_cluster.append(r)
                else:
                    break   # remaining thick rules are mid-table markers

            if len(header_cluster) < 2:
                continue

            # Each consecutive pair in the header cluster = one header band
            header_bands: list[tuple[float, float]] = []
            for i in range(len(header_cluster) - 1):
                header_bands.append((header_cluster[i]["top"], header_cluster[i + 1]["bottom"]))

            table_top     = header_cluster[0]["top"]
            header_bottom = header_cluster[-1]["bottom"]

            # Mid-table thick rules (beyond the header cluster) become extra row
            # separators so their content is extracted as bold repeat-header rows.
            mid_thick = thick[len(header_cluster):]
            mid_thick_tops = [r["top"] for r in mid_thick] + [r["bottom"] for r in mid_thick]

            row_seps = sorted(
                [r["top"] for r in thin if r["top"] > header_bottom] + mid_thick_tops
            )
            if not row_seps:
                continue

            table_bottom = thin[-1]["bottom"] if thin else mid_thick[-1]["bottom"] if mid_thick else header_bottom + 20

            t_words = [
                w for w in words
                if w["x0"] >= x0 - 5 and w["x1"] <= x1 + 5
                and w["top"] >= table_top and w["top"] <= table_bottom + 10
            ]
            if not t_words:
                continue

            # Infer column boundaries — use the header band with the most
            # distinct X start positions (most complete column definition).
            # This handles cases where the last header band is a straddle row
            # with fewer columns than the actual column-header row.
            best_col_words = None
            best_x_count = 0
            for hdr_top_c, hdr_bot_c in header_bands:
                cw = [w for w in t_words
                      if hdr_top_c - 1 <= w["top"] <= hdr_bot_c - 1]
                # Count distinct X positions with meaningful gaps
                xs = sorted(set(round(w["x0"], 0) for w in cw))
                distinct = 1
                for xi in range(1, len(xs)):
                    if xs[xi] - xs[xi-1] > _ROW_SHOW_COL_GAP:
                        distinct += 1
                if distinct > best_x_count:
                    best_x_count = distinct
                    best_col_words = cw
            if not best_col_words:
                best_col_words = [w for w in t_words if w["top"] >= thick[0]["bottom"]]
            if not best_col_words:
                best_col_words = t_words

            last_hdr_top, last_hdr_bot = header_bands[-1]

            # Use whitespace projection on HEADER words only.  Using all table
            # words (including data rows) creates many 10-12pt inter-word gaps
            # across the row that are indistinguishable from real column gaps,
            # causing every inter-word space to become a false column break.
            # Header words have one cluster per column with clear gaps between
            # clusters, making them far more reliable for column detection.
            # Fall back to all words (with a larger min_gap guard) only when the
            # header is so sparse that a single-column result would be returned.
            col_breaks = _col_breaks_from_projection(best_col_words, x0, x1)
            if len(col_breaks) <= 2:
                # Sparse header: retry with all table words and a tighter gap
                # guard (but still > typical inter-word space of ~2pt)
                col_breaks = _col_breaks_from_projection(t_words, x0, x1)
            n_cols = len(col_breaks) - 1

            def _assign_col(wx: float, _col_breaks=col_breaks, _n_cols=n_cols) -> int:
                # Use 2pt left tolerance to handle float rounding between
                # word x0 (raw PDF units) and col_breaks (rounded rect x0).
                for ci in range(len(_col_breaks) - 1):
                    if _col_breaks[ci] - 2 <= wx < _col_breaks[ci + 1]:
                        return ci
                return _n_cols - 1

            def _words_in_band(top_y: float, bot_y: float, tight: bool = False,
                               _t_words=t_words, _n_cols=n_cols) -> list[str]:
                tol = 1 if tight else 2
                band = [w for w in _t_words
                        if w["top"] >= top_y - tol and w["top"] <= bot_y - tol]
                cells = [""] * _n_cols
                cell_bold = [False] * _n_cols
                cell_last_word_idx = [-1] * _n_cols  # track last word position per cell
                for w in band:
                    tm = _tm_type(w)
                    c = _assign_col(w["x0"])
                    if tm is not None:
                        # Append TM sentinel to the last text in this cell
                        if cells[c]:
                            cells[c] = _encode_tm(cells[c], tm)
                    else:
                        cells[c] = (cells[c] + " " + w["text"]).strip()
                        if "Bold" in w.get("fontname", ""):
                            cell_bold[c] = True
                result = []
                for i, cell in enumerate(cells):
                    if cell_bold[i] and cell:
                        result.append(f"__BOLD__{cell}")
                    else:
                        result.append(cell)
                return result

            rows: list[list[str]] = []

            # Header rows — detect straddled (spanning) cells
            # Strategy: if the leftmost word in a header band starts well to the
            # right of col1's left edge, treat it as a spanning (straddle) cell.
            col1_right_threshold = col_breaks[0] + (col_breaks[1] - col_breaks[0]) * 0.5

            for hdr_top, hdr_bot in header_bands:
                hdr = _words_in_band(hdr_top, hdr_bot, tight=True)
                non_empty = [i for i, c in enumerate(hdr) if c.strip()]

                band_words = [w for w in t_words
                              if w["top"] >= hdr_top - 1 and w["top"] <= hdr_bot - 1]
                leftmost_x = min((w["x0"] for w in band_words), default=x0)

                is_straddle = (
                    n_cols > 1
                    and leftmost_x > col1_right_threshold
                    and len(non_empty) > 0
                )

                if is_straddle:
                    all_text = " ".join(w["text"] for w in
                                        sorted(band_words, key=lambda w: w["x0"]))
                    straddle_row = [""] * n_cols
                    straddle_row[0] = all_text.strip()
                    straddle_row[1] = f"__STRADDLE__{n_cols}"
                    rows.append(straddle_row)
                else:
                    rows.append(hdr)

            # Data rows
            band_tops = [header_bottom] + row_seps
            band_bots = row_seps + [table_bottom + 15]
            for top_y, bot_y in zip(band_tops, band_bots):
                row = _words_in_band(top_y, bot_y)
                if any(c.strip() for c in row):
                    rows.append(row)

            results.append((rows, table_top, table_bottom))

    return results


def _parse_page_range(page_range: str, total_pages: int) -> set[int]:
    """
    Parse a page range string like "1-5, 8, 12-15" into a set of
    0-based page indices. Returns None if page_range is empty (= all pages).

    Examples:
        "1-5, 8"     → {0, 1, 2, 3, 4, 7}
        "3"          → {2}
        ""           → None (all pages)
    """
    if not page_range or not page_range.strip():
        return None
    indices = set()
    for part in page_range.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            start_s, end_s = part.split("-", 1)
            start = max(1, int(start_s.strip()))
            end   = min(total_pages, int(end_s.strip()))
            for i in range(start, end + 1):
                indices.add(i - 1)  # convert to 0-based
        else:
            page_num = int(part)
            if 1 <= page_num <= total_pages:
                indices.add(page_num - 1)
    return indices if indices else None


_BLANK_PAGE_PATTERNS = [
    re.compile(r"^this\s+page\s+(is\s+)?intentionally\s+(left\s+)?blank", re.IGNORECASE),
    re.compile(r"^intentionally\s+(left\s+)?blank", re.IGNORECASE),
    re.compile(r"^this\s+page\s+left\s+blank", re.IGNORECASE),
    re.compile(r"^blank\s+page$", re.IGNORECASE),
]


def _should_drop(text: str) -> bool:
    t = text.strip()
    if not t:
        return True
    if len(t) < 3:
        return True
    for pat in _DROP_PATTERNS:
        if pat.search(t):
            return True
    return False


def _is_blank_page(page_text: str) -> bool:
    """Return True if the page contains only a blank-page notice or nothing."""
    cleaned = page_text.strip()
    if not cleaned:
        return True
    lines = [l.strip() for l in cleaned.splitlines() if l.strip()]
    meaningful = [l for l in lines if not _should_drop(l)]
    if not meaningful:
        return True
    full = " ".join(meaningful)
    for pat in _BLANK_PAGE_PATTERNS:
        if pat.match(full):
            return True
    return False


# ---------------------------------------------------------------------------
# PDF Extractor
# ---------------------------------------------------------------------------

# Font-size → heading level mapping (calibrated on Gilbarco manuals)
_H1_SIZE   = 17.0   # ≥ 17 pt bold → H1
_H2_SIZE   = 13.5   # ≥ 13.5 pt bold → H2
_H3_SIZE   = 11.5   # ≥ 11.5 pt bold → H3
_NOTE_SIZE = 14.0   # ≥ 14 pt bold → potential note header
_STEP_SIZE =  9.5   # ≤ 9.5 pt → running header/footer (drop)
_DROP_SIZE =  9.5


def _classify_line(word_group: list[dict]) -> tuple[str, int]:
    """Return (block_type, level) for a group of words on one line."""
    if not word_group:
        return "paragraph", 0

    sizes = [w.get("size", 11) for w in word_group]
    avg_size = sum(sizes) / len(sizes)

    fonts = [w.get("fontname", "") for w in word_group]
    is_bold = any("Bold" in f or "BoldMT" in f for f in fonts)

    if avg_size <= _DROP_SIZE:
        return "dropped", 0

    if is_bold:
        if avg_size >= _H1_SIZE:
            return "heading", 1
        if avg_size >= _NOTE_SIZE:
            return "note_header", 0
        if avg_size >= _H2_SIZE:
            return "heading", 2
        if avg_size >= _H3_SIZE:
            return "heading", 3

    return "paragraph", 0



# ---------------------------------------------------------------------------
# Trademark / superscript detection
# ---------------------------------------------------------------------------
_TM_SYMBOLS = {
    "®": "reg",
    "®": "reg",
    "™": "tm",
    "™": "tm",
    "SM": "service",   # only when superscript (size <= 6)
}
_TM_SUPERSCRIPT_SIZE = 6.0  # pt — words at this size or smaller are superscripts


def _tm_type(word: dict) -> str | None:
    """Return tm type if this word is a trademark superscript, else None."""
    text = word.get("text", "").strip()
    size = word.get("size", 99)
    if size > _TM_SUPERSCRIPT_SIZE:
        return None
    return _TM_SYMBOLS.get(text)


def _encode_tm(text: str, tm_type: str) -> str:
    """Append a __TM__{type}__ sentinel to the preceding text."""
    return text.rstrip() + f"__TM__{tm_type}__"

def extract_pdf(file_bytes: bytes, page_range: str = "") -> list[dict]:
    """Extract a content tree from a text-based PDF.

    Args:
        file_bytes:  Raw bytes of the PDF file.
        page_range:  Optional page range string e.g. "1-5, 8, 12-15".
                     Leave empty to extract all pages.
    """
    import pdfplumber  # type: ignore

    blocks: list[dict] = []
    dropped_count = 0
    blank_pages_skipped = 0

    with pdfplumber.open(file_bytes if hasattr(file_bytes, "read") else
                         __import__("io").BytesIO(file_bytes)) as pdf:

        total_pages = len(pdf.pages)
        total_chars = sum(len(p.extract_text() or "") for p in pdf.pages)
        if total_chars < 50:
            raise ExtractorError(
                "No extractable text found. This appears to be a scanned PDF. "
                "Please supply a text-based (digital) PDF."
            )

        # Resolve page range filter
        page_indices = _parse_page_range(page_range, total_pages)  # None = all pages

        for page_idx, page in enumerate(pdf.pages):

            # ---- Page range filter (B-001) ----
            if page_indices is not None and page_idx not in page_indices:
                continue

            # ---- Blank page detection (B-002) ----
            page_text = page.extract_text() or ""
            if _is_blank_page(page_text):
                blank_pages_skipped += 1
                continue

            # ---- Tables: pdfplumber bordered + ROW_SHOW borderless ----
            # We collect page-level blocks with Y positions so we can
            # reorder them correctly (tables are extracted before words by
            # pdfplumber, which would put them before any lead-in paragraph).
            page_blocks: list[tuple[float, dict]] = []  # (y_pos, block)

            # Pass 1: pdfplumber's standard table detector (bordered tables)
            std_tables   = page.extract_tables()
            std_bboxes: list = []
            std_table_objs = []
            if std_tables:
                std_table_objs = page.find_tables()
                std_bboxes = [t.bbox for t in std_table_objs]

            for ti, table_data in enumerate(std_tables):
                if not table_data:
                    continue
                rows = [[cell or "" for cell in row] for row in table_data]
                y_pos = std_bboxes[ti][1] if ti < len(std_bboxes) else 0.0
                page_blocks.append((y_pos, make_block("table", "", is_header=True, rows=rows)))

            # Pass 2: ROW_SHOW borderless table detector — runs BEFORE Pass 1.5
            # so that its (correct) column detection takes priority over pdfplumber's
            # text-strategy column detection, which fragments multi-word column headers.
            rs_bboxes: list = []
            for rs_rows, rs_y, rs_ybot in _extract_rowshow_tables(page):
                if not rs_rows:
                    continue
                # Skip if Y overlaps a Pass-1 bordered table
                skip = False
                for bbox in std_bboxes:
                    if len(bbox) == 4:
                        _, btop, _, bbot = bbox
                        if btop - 20 <= rs_y <= bbot + 20:
                            skip = True
                            break
                if not skip:
                    blk = make_block("table", "", is_header=True, rows=rs_rows)
                    blk["metadata"]["n_header_rows"] = 1
                    page_blocks.append((rs_y, blk))
                    # Record bbox so Pass 1.5 can skip this region
                    rs_bboxes.append((0, rs_y, 999, rs_ybot))

            # Merge ROW_SHOW bboxes into std_bboxes before Pass 1.5 runs
            std_bboxes = std_bboxes + rs_bboxes

            # Pass 1.5: pdfplumber text-strategy detector (horizontal rules +
            # text-aligned columns).  Only used for tables missed by both Pass 1
            # and Pass 2 (ROW_SHOW).  Running after ROW_SHOW ensures that tables
            # already captured with correct column layout are not re-detected with
            # pdfplumber's text-strategy column detection, which fragments multi-word
            # column headers into many narrow false columns.
            _TEXT_STRAT = {
                "vertical_strategy":    "text",
                "horizontal_strategy":  "lines",
                "snap_tolerance":       3,
                "join_tolerance":       3,
                "min_words_vertical":   3,
                "min_words_horizontal": 2,
                "intersection_tolerance": 3,
            }
            try:
                ts_table_objs = page.find_tables(table_settings=_TEXT_STRAT)
            except Exception:
                ts_table_objs = []
            ts_added_bboxes: list = []

            page_height = page.height or 792.0
            for ti, tbl in enumerate(ts_table_objs):
                bbox = tbl.bbox
                _, btop, _, bbot = bbox
                # Reject tables that span >70% of the page height — these are
                # almost always false positives caused by page margin lines being
                # treated as table row separators, turning the whole page into one
                # giant "table".
                if (bbot - btop) / page_height > 0.70:
                    continue
                # Skip if Y overlaps Pass-1 bordered table OR Pass-2 ROW_SHOW table
                overlaps = any(
                    len(b) == 4 and b[1] - 20 <= btop <= b[3] + 20
                    for b in std_bboxes
                )
                if overlaps:
                    continue
                raw = tbl.extract()
                if not raw:
                    continue
                rows = [[str(c or "").strip() for c in row] for row in raw]
                if len(rows) < 2:
                    continue
                if any(any(cell for cell in row) for row in rows):
                    blk = make_block("table", "", is_header=True, rows=rows)
                    blk["metadata"]["n_header_rows"] = 1
                    page_blocks.append((btop, blk))
                    ts_added_bboxes.append(tbl.bbox)

            std_bboxes = std_bboxes + ts_added_bboxes

            # ---- Words → lines ----
            words = page.extract_words(
                x_tolerance=3,
                y_tolerance=3,
                keep_blank_chars=False,
                use_text_flow=True,
                extra_attrs=["fontname", "size"],
            )

            # Build set of Y ranges to exclude — areas covered by extracted tables
            excluded_y_bands: list[tuple[float, float]] = []
            for bbox in std_bboxes:
                if len(bbox) == 4:
                    excluded_y_bands.append((bbox[1] - 2, bbox[3] + 2))

            def _in_table_area(y: float) -> bool:
                for y0, y1 in excluded_y_bands:
                    if y0 <= y <= y1:
                        return True
                return False

            # Group words into lines by top-coordinate
            # Superscript TM symbols (size <= 6) are merged into the
            # preceding line rather than creating their own line.
            lines: dict[float, list] = {}
            for w in words:
                top = round(w["top"], 1)
                if _in_table_area(top):
                    continue
                tm = _tm_type(w)
                if tm is not None:
                    # Find the most recent line above this superscript (within 8pt)
                    parent_top = None
                    for lt in sorted(lines.keys(), reverse=True):
                        if abs(top - lt) <= 8:
                            parent_top = lt
                            break
                    if parent_top is not None:
                        # Tag the last word in that line with TM sentinel
                        parent_line = lines[parent_top]
                        if parent_line:
                            last = parent_line[-1]
                            # Create a synthetic word with TM-encoded text
                            tagged = dict(last)
                            tagged["text"] = _encode_tm(last["text"], tm)
                            parent_line[-1] = tagged
                        continue  # don't add superscript as its own word
                lines.setdefault(top, []).append(w)

            prev_para = None
            for top in sorted(lines):
                word_group = lines[top]
                text = " ".join(w["text"] for w in word_group).strip()
                # Detect if the entire line is bold (all words have Bold in fontname)
                line_is_bold = (
                    len(word_group) > 0
                    and all("Bold" in w.get("fontname", "") for w in word_group)
                )

                if _should_drop(text):
                    dropped_count += 1
                    continue

                block_type, level = _classify_line(word_group)

                if block_type == "dropped":
                    dropped_count += 1
                    continue

                # Bullet detection
                if text.startswith(("•", "–", "-", "▪", "◆")) or \
                   re.match(r"^[●○■□▸▹►]", text):
                    text = re.sub(r"^[•–\-▪◆●○■□▸▹►]\s*", "", text)
                    block_type = "list_item"
                    meta = {"list_kind": "bullet"}
                    page_blocks.append((top, make_block(block_type, text, metadata=meta)))
                    prev_para = None
                    continue

                # Numbered item detection
                num_match = re.match(r"^(\d{1,2})\s{1,4}(.+)", text)
                if num_match and block_type == "paragraph":
                    text = num_match.group(2)
                    block_type = "list_item"
                    meta = {"list_kind": "numbered", "num": int(num_match.group(1))}
                    page_blocks.append((top, make_block(block_type, text, metadata=meta)))
                    prev_para = None
                    continue

                # Figure caption
                if re.match(r"^Figure\s+\d+\s*:", text, re.IGNORECASE):
                    page_blocks.append((top, make_block("figure", text)))
                    prev_para = None
                    continue

                # Inline note
                if re.match(r"^Notes?:", text, re.IGNORECASE):
                    page_blocks.append((top, make_block("note_inline", text)))
                    prev_para = None
                    continue

                # Code block signals
                code_signals = ("telnet ", "C:\\>", "$ ", "http://")
                if any(text.startswith(s) for s in code_signals):
                    page_blocks.append((top, make_block("code_block", text)))
                    prev_para = None
                    continue

                # Paragraph merging (continuation lines at same style)
                if block_type == "paragraph" and prev_para is not None and page_blocks:
                    last = page_blocks[-1][1]
                    if last["type"] == "paragraph" and not last["text"].endswith((".", ":", "?")):
                        last["text"] = last["text"] + " " + text
                        continue

                blk = make_block(block_type, text, level=level)
                if line_is_bold and block_type == "paragraph":
                    blk["metadata"]["bold"] = True
                page_blocks.append((top, blk))
                prev_para = block_type if block_type == "paragraph" else None

            # ---- Flush page blocks in Y order ----
            page_blocks.sort(key=lambda x: x[0])
            blocks.extend(b for _, b in page_blocks)

    # Tag how many blocks were dropped
    for b in blocks:
        b.setdefault("metadata", {})
    if blocks:
        blocks[0]["metadata"]["dropped_count"]       = dropped_count
        blocks[0]["metadata"]["blank_pages_skipped"] = blank_pages_skipped

    return blocks


# ---------------------------------------------------------------------------
# DOCX Extractor
# ---------------------------------------------------------------------------

_DOCX_STYLE_MAP = {
    "Heading 1": ("heading", 1),
    "Heading 2": ("heading", 2),
    "Heading 3": ("heading", 3),
    "Heading 4": ("heading", 3),
    "Title":     ("heading", 1),
    "Subtitle":  ("heading", 2),
}

_DOCX_NOTE_STYLES = {"Caution", "Warning", "Note", "Important"}


def extract_docx(file_bytes: bytes, image_folder: str = "") -> list[dict]:
    """Extract a content tree from a DOCX file.

    Args:
        file_bytes: Raw bytes of the .docx file.
        image_folder: Optional path to the extracted media folder (from the
                      renamed .zip). When provided, image relationships are
                      resolved to absolute paths for DITA <image href>.
    """
    import io
    from docx import Document  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    doc = Document(io.BytesIO(file_bytes))
    blocks: list[dict] = []
    dropped_count = 0
    image_map: dict[str, str] = {}

    # Build image relationship map if folder provided
    if image_folder:
        img_dir = Path(image_folder)
        if img_dir.is_dir():
            # Map rId → absolute path by scanning rels
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    target = rel.target_ref  # e.g. "media/image1.png"
                    fname = Path(target).name
                    candidate = img_dir / fname
                    if candidate.exists():
                        image_map[rel.rId] = str(candidate)

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        if not text:
            continue

        if _should_drop(text):
            dropped_count += 1
            continue

        # Check for inline images in runs
        for run in para.runs:
            for drawing in run._element.findall(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"):
                blip_fills = drawing.findall(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/picture}blipFill")
                for bf in blip_fills:
                    blip = bf.find(
                        "{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
                    if blip is not None:
                        r_embed = blip.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        img_path = image_map.get(r_embed, "")
                        caption = text or f"Image {r_embed}"
                        blocks.append(make_block(
                            "figure", caption,
                            metadata={"image_href": img_path, "r_id": r_embed}
                        ))

        # Style-driven classification
        if style_name in _DOCX_STYLE_MAP:
            btype, level = _DOCX_STYLE_MAP[style_name]
            blocks.append(make_block(btype, text, level=level))
            continue

        # Note styles
        if style_name in _DOCX_NOTE_STYLES:
            blocks.append(make_block("note_header", text))
            continue

        # List paragraph
        if "List" in style_name:
            list_kind = "numbered" if "Number" in style_name else "bullet"
            blocks.append(make_block("list_item", text, metadata={"list_kind": list_kind}))
            continue

        # Inline note prefix
        if re.match(r"^Notes?:", text, re.IGNORECASE):
            blocks.append(make_block("note_inline", text))
            continue

        # Figure caption
        if re.match(r"^Figure\s+\d+\s*:", text, re.IGNORECASE):
            blocks.append(make_block("figure", text))
            continue

        # Code style
        if "Code" in style_name or "Preformatted" in style_name:
            blocks.append(make_block("code_block", text))
            continue

        # Bullet by text prefix
        if text.startswith(("•", "–", "▪")):
            text = re.sub(r"^[•–▪]\s*", "", text)
            blocks.append(make_block("list_item", text, metadata={"list_kind": "bullet"}))
            continue

        # Numbered item
        num_match = re.match(r"^(\d{1,2})[.)]\s+(.+)", text)
        if num_match:
            blocks.append(make_block(
                "list_item", num_match.group(2),
                metadata={"list_kind": "numbered", "num": int(num_match.group(1))}
            ))
            continue

        blocks.append(make_block("paragraph", text))

    # Tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            blocks.append(make_block("table", "", is_header=True, rows=rows))

    if blocks:
        blocks[0]["metadata"]["dropped_count"] = dropped_count

    return blocks
